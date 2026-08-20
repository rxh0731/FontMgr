"""经文正文只读提取服务回归测试。"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document

from services.scripture_text_service import (
    load_scripture_text,
    scripture_document_filter,
)


class ScriptureTextServiceTests(unittest.TestCase):
    def test_filter_lists_common_text_and_office_formats(self) -> None:
        document_filter = scripture_document_filter()

        for extension in ("*.txt", "*.doc", "*.docx", "*.pdf", "*.wps", "*.xlsx", "*.pptx"):
            self.assertIn(extension, document_filter)

    def test_loads_gb18030_plain_text_without_modifying_source(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "经文.txt"
            raw = "般若波罗蜜多心经\n观自在菩萨".encode("gb18030")
            path.write_bytes(raw)

            result = load_scripture_text(str(path))

            self.assertEqual(result.text, "般若波罗蜜多心经\n观自在菩萨")
            self.assertEqual(result.source_name, "经文.txt")
            self.assertEqual(path.read_bytes(), raw)

    def test_plain_text_keeps_leading_and_trailing_cell_spaces(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "占格空格.txt"
            path.write_text("  甲乙　\n\n丙丁  ", encoding="utf-8")

            result = load_scripture_text(str(path))

            self.assertEqual(result.text, "  甲乙　\n\n丙丁  ")

    def test_loads_docx_paragraphs_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "经文.docx"
            document = Document()
            document.add_paragraph("第一段")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "第二段"
            table.cell(0, 1).text = "第三段"
            document.save(path)

            result = load_scripture_text(str(path))

            self.assertIn("第一段", result.text)
            self.assertIn("第二段\t第三段", result.text)

    def test_docx_keeps_paragraph_and_cell_boundary_spaces(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "占格空格.docx"
            document = Document()
            document.add_paragraph("  第一段　")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = " 第二段"
            table.cell(0, 1).text = "第三段 "
            document.save(path)

            result = load_scripture_text(str(path))

            self.assertIn("  第一段　", result.text)
            self.assertIn(" 第二段\t第三段 ", result.text)

    def test_loads_pdf_text_layer(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scripture.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Scripture text")
            document.save(path)
            document.close()

            result = load_scripture_text(str(path))

            self.assertIn("Scripture text", result.text)

    def test_rejects_unknown_format(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "经文.bin"
            path.write_bytes(b"text")

            with self.assertRaisesRegex(ValueError, "暂不支持"):
                load_scripture_text(str(path))


if __name__ == "__main__":
    unittest.main()
